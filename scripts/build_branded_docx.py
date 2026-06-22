#!/usr/bin/env python3
"""Build a house-style-branded ``DESIGN_DOC.docx`` from a reference template.

The *content* is exactly ``docs/DESIGN_DOC.md`` (same data as the plain
``npm run docs:docx`` build). The *look and feel* — cover page, front-matter
control tables, embedded fonts, theme colours, the confidential footer with
page numbers, and the running header — all come from a **reference .docx you
supply** (a branded Word document in your organisation's / your client's house
style). The reference document is never committed (it carries third-party
branding and licensed embedded fonts); pass its path with ``--template`` or the
``BRANDED_TEMPLATE`` env var.

Everything site-specific (the document title/author/date, the cover/header
search-replace strings, and the three control-table contents) is read from an
optional JSON config so this script itself names no client and ships clean.
See ``scripts/branded_docx.config.example.json`` for the shape; put your real
values in ``scripts/branded_docx.config.json`` (git-ignored) or pass
``--config <path>``. Metadata fields also accept ``DOC_TITLE`` / ``DOC_AUTHOR``
/ … env overrides.

How it works
------------
1. A build-time copy of the markdown is produced with the leading title block
   and the inline "Document Control" table removed — that information moves onto
   the cover page and the three front-matter tables, so we don't duplicate it.

2. ``pandoc --reference-doc=<template>`` renders the body. Because pandoc
   inherits the reference document's styles, theme, embedded fonts and every
   header/footer part (keeping their relationship IDs), the rendered body is
   *already* in the template's visual system. Pandoc also drops the template's
   heading auto-numbering, so our manual section numbers ("1.", "2." …) are not
   doubled.

3. We splice the template cover page + the three control tables + the
   front-matter section break to the top of the body, then restore the
   template's body section properties (which carry the confidential footer).

4. Finally we drop our document's text onto the cover (title / subtitle / date),
   populate the three control tables, and refresh the stale running header —
   all driven by the config's search-replace maps so the surgery stays matched
   to whichever reference template you supply.

Usage
-----
    python3 scripts/build_branded_docx.py [--template <ref.docx>] \
        [--config <config.json>] [--out <out.docx>]
"""
import argparse
import copy
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile

from lxml import etree
from docx import Document
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn

XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

# ---------------------------------------------------------------------------
# Neutral defaults. Everything here is overridable via the JSON config (and the
# metadata fields via DOC_* env vars) so nothing client-specific is committed.
# The search-replace maps default to empty because their "needles" must match
# the exact placeholder text inside YOUR reference template — see the example
# config for how to populate them.
# ---------------------------------------------------------------------------
DEFAULT_META = {
    "title": "AEP Data Lifecycle Helper",
    "subtitle": "Design & Architecture Document",
    "date": "June 2026",
    "author": "Adobe",
    "subject": "AEP Data Lifecycle Helper design and architecture",
    "keywords": "Confidential",
}

# Neutral control-table defaults (token-substituted at load). Override in config.
DEFAULT_DOCUMENT_CONTROL = [
    ["Document Title", "{title} — {subtitle}"],
    ["Authors", "{author}"],
    ["Document Purpose",
     "Design and architecture reference for the AEP Data Lifecycle Helper — a "
     "local, operator-run tool for quota-safe bulk identity deletion via the "
     "Adobe Experience Platform Data Hygiene API."],
    ["Status", "Production-ready"],
    ["Classification", "Confidential"],
    ["Intended Audience", "Platform architects, security reviewers, operators"],
]
DEFAULT_VERSION_CONTROL = [
    ["1.0.0", "{date}", "{author}", "Initial release."],
]
DEFAULT_REFERENCES = [
    ["1", "—", "{date}", "{author}", "ARCHITECTURE.md — living system overview"],
    ["2", "—", "{date}", "{author}", "CHANGELOG.md — full change history"],
    ["3", "—", "{date}", "{author}", "REVIEW.md — Adobe API payload reference"],
]

ENV_META_KEYS = {
    "title": "DOC_TITLE", "subtitle": "DOC_SUBTITLE", "date": "DOC_DATE",
    "author": "DOC_AUTHOR", "subject": "DOC_SUBJECT", "keywords": "DOC_KEYWORDS",
}

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
# Namespaces the cover-page drawings use that the template root may declare only
# locally; declare them on our parse wrapper so lxml never trips on a prefix.
EXTRA_NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "a14": "http://schemas.microsoft.com/office/drawing/2010/main",
    "a16": "http://schemas.microsoft.com/office/drawing/2014/main",
}


def w(tag):
    return f"{{{W_NS}}}{tag}"


def _esc(s):
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _subst(value, meta, xml=False):
    """Expand {title}/{subtitle}/{date}/{author} tokens in a config value.

    For values spliced into raw XML (cover/header/footer needles) pass
    xml=True so the expansion is XML-escaped; for plain table-cell text the
    raw metadata is used (python-docx escapes on write)."""
    out = value
    for tok, key in (("{title}", "title"), ("{subtitle}", "subtitle"),
                     ("{date}", "date"), ("{author}", "author")):
        repl = _esc(meta[key]) if xml else meta[key]
        out = out.replace(tok, repl)
    return out


def load_site(config_path):
    """Resolve all site-specific text: DEFAULTS < JSON config file < DOC_* env.

    Returns a dict with metadata plus the resolved (token-substituted)
    search-replace maps and control-table rows. The config file is optional;
    without it the build runs with neutral defaults (and leaves the template's
    own cover placeholders in place — a warning is printed)."""
    cfg = {}
    if config_path and os.path.isfile(config_path):
        try:
            with open(config_path, encoding="utf8") as fh:
                cfg = json.load(fh)
        except (OSError, ValueError) as e:
            raise SystemExit(f"build_branded_docx: could not read config "
                             f"{config_path!r}: {e}")

    meta = dict(DEFAULT_META)
    for key, env in ENV_META_KEYS.items():
        if cfg.get(key):
            meta[key] = cfg[key]
        if os.environ.get(env):                       # env wins over config
            meta[key] = os.environ[env]

    def _map(name):
        return {needle: _subst(val, meta, xml=True)
                for needle, val in (cfg.get(name) or {}).items()}

    def _rows(name, default):
        rows = cfg.get(name) or default
        return [[_subst(str(c), meta, xml=False) for c in row] for row in rows]

    return {
        "meta": meta,
        "cover": _map("coverReplacements"),
        "chrome": _map("chromeReplacements"),
        "footer": _map("footerReplacements"),
        "documentControl": _rows("documentControl", DEFAULT_DOCUMENT_CONTROL),
        "versionControl": _rows("versionControl", DEFAULT_VERSION_CONTROL),
        "references": _rows("references", DEFAULT_REFERENCES),
    }


def strip_front_matter(md_text):
    """Drop everything before the Diagrams blockquote (title + Document Control
    move to the cover page / control tables), and drop the manual Table of
    Contents (the front matter carries an auto-updating Word TOC field)."""
    marker = "> **Diagrams & screenshots** in this document:"
    idx = md_text.find(marker)
    if idx == -1:
        raise SystemExit("build_branded_docx: could not find the Diagrams "
                         "blockquote marker in DESIGN_DOC.md — strip boundary "
                         "changed?")
    body = md_text[idx:]

    toc = body.find("**Table of Contents**")
    if toc != -1:
        nxt = body.find("\n## ", toc)
        if nxt == -1:
            raise SystemExit("build_branded_docx: could not find the first "
                             "section after the manual Table of Contents.")
        body = body[:toc] + body[nxt + 1:]   # drop TOC heading + list + the '---'
    return body


def run_pandoc(body_md_path, template, out_path):
    # The body's top-level sections are markdown "##" (the document title "#" is
    # stripped to the cover). Shift up one level so sections render as the
    # template's Heading 1 — matching house styles where numbered sections are
    # Heading 1. Safe: the body has no real H1 (so nothing collapses to a part
    # header), and the template's heading auto-numbering is inert (pandoc drops
    # it), so our manual "1.", "2." numbers are not doubled.
    cmd = [
        "pandoc", "-f", "markdown", "--resource-path=docs",
        "--shift-heading-level-by=-1",
        body_md_path, "--reference-doc", template, "-o", out_path,
    ]
    subprocess.run(cmd, check=True)


def extract_template_parts(template):
    """Return (front_matter_xml, body_sectPr_xml, document_root_tag)."""
    with zipfile.ZipFile(template) as z:
        doc = z.read("word/document.xml").decode("utf8")

    root_tag = re.search(r"<w:document\b[^>]*>", doc).group(0)
    # Ensure the drawing namespaces are declared on the root we reuse for parsing.
    for prefix, uri in EXTRA_NS.items():
        if f"xmlns:{prefix}=" not in root_tag:
            root_tag = root_tag[:-1] + f' xmlns:{prefix}="{uri}">'

    body_open = doc.find("<w:body>") + len("<w:body>")
    # The first <w:sectPr> is the front-matter section break (lives inside a
    # paragraph's pPr). Front matter = body start .. end of that paragraph.
    s1 = doc.find("<w:sectPr")
    s1_close = doc.find("</w:sectPr>", s1) + len("</w:sectPr>")
    front_end = doc.find("</w:p>", s1_close) + len("</w:p>")
    front = doc[body_open:front_end]

    # The last <w:sectPr> is the body section properties (direct child of body).
    s2 = doc.rfind("<w:sectPr")
    s2_close = doc.find("</w:sectPr>", s2) + len("</w:sectPr>")
    body_sect = doc[s2:s2_close]

    return front, body_sect, root_tag


def edit_cover_text(front_xml, cover_replacements):
    """Replace the template's placeholder cover title / subtitle / date using
    the config's coverReplacements map (needle -> value)."""
    if not cover_replacements:
        print("build_branded_docx: WARNING — no coverReplacements in config; "
              "the template's own cover placeholder text will be left as-is.",
              file=sys.stderr)
    for needle, value in cover_replacements.items():
        front_xml = front_xml.replace(needle, value)
    return front_xml


def parse_fragment(root_tag, inner_xml):
    """Parse a body-fragment string into a list of top-level elements.

    Uses python-docx's parser so the elements get the registered oxml classes
    (CT_Tbl/CT_P/...), which is what makes ``doc.tables`` work on them after
    insertion. The full root tag declares every namespace the fragment uses.
    """
    wrapper = f"{root_tag}<w:body>{inner_xml}</w:body></w:document>"
    tree = parse_xml(wrapper)
    return list(tree.find(w("body")))


def parse_sectpr(root_tag, sect_xml):
    wrapper = f"{root_tag}<w:body>{sect_xml}</w:body></w:document>"
    tree = parse_xml(wrapper)
    return tree.find(w("body")).find(w("sectPr"))


def build_toc_elements(root_tag):
    """A clean, auto-updating Word TOC field (Heading 1-3) to replace the
    template's stale cached TOC. The 'TOC Heading' style is excluded from the
    TOC itself; w:dirty + settings/updateFields make Word rebuild it on open."""
    toc_xml = (
        '<w:p><w:pPr><w:pStyle w:val="TOCHeading"/></w:pPr>'
        '<w:r><w:t>Table of Contents</w:t></w:r></w:p>'
        '<w:p>'
        '<w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        '<w:r><w:rPr><w:i/><w:color w:val="6B737B"/></w:rPr>'
        '<w:t xml:space="preserve">Right-click here and choose “Update Field” '
        '(or select the document and press F9) to build the table of contents.</w:t></w:r>'
        '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
        '</w:p>'
    )
    return parse_fragment(root_tag, toc_xml)


def swap_stale_toc(front_children, root_tag):
    """Replace the template's stale TOC content-control (the <w:sdt> whose field
    code is a TOC) with a fresh, empty TOC field."""
    out, replaced = [], False
    for el in front_children:
        is_toc_sdt = (
            etree.QName(el).localname == "sdt"
            and any("TOC" in (t.text or "") for t in el.iter(w("instrText")))
        )
        if is_toc_sdt and not replaced:
            out.extend(build_toc_elements(root_tag))
            replaced = True
        else:
            out.append(el)
    if not replaced:
        raise SystemExit("build_branded_docx: stale TOC content-control not "
                         "found in the template front matter — structure changed?")
    return out


def set_cell_text(cell, text):
    """Replace a table cell's content with a single run of ``text``.

    Some control-table cells wrap their value in property-bound content
    controls (<w:sdt> bound to Title/Subject) or span multiple paragraphs;
    editing only the visible runs would leave that content behind (and Word
    would re-populate the bound controls on field update). So we clear the cell
    down to its <w:tcPr> and rebuild one clean paragraph, carrying over the
    original paragraph/run formatting so the cell styling is preserved."""
    tc = cell._tc
    pPr = rPr = None
    first_p = tc.find(qn("w:p"))
    if first_p is not None and first_p.find(qn("w:pPr")) is not None:
        pPr = copy.deepcopy(first_p.find(qn("w:pPr")))
    for r in tc.iter(qn("w:r")):                       # first run anywhere in cell
        if r.find(qn("w:rPr")) is not None:
            rPr = copy.deepcopy(r.find(qn("w:rPr")))
            break

    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)

    p = OxmlElement("w:p")
    if pPr is not None:
        p.append(pPr)
    run = OxmlElement("w:r")
    if rPr is not None:
        run.append(rPr)
    t = OxmlElement("w:t")
    t.set(XML_SPACE, "preserve")
    t.text = text
    run.append(t)
    p.append(run)
    tc.append(p)


def ensure_rows(table, n_data_rows, first_data_idx):
    """Grow/shrink a table so it has exactly first_data_idx + n_data_rows rows,
    cloning the template data row (so cloned rows keep cell widths/borders)."""
    tbl = table._tbl
    target = first_data_idx + n_data_rows
    template_tr = table.rows[first_data_idx]._tr
    while len(table.rows) < target:
        tbl.append(copy.deepcopy(template_tr))
    # Remove surplus trailing rows.
    while len(table.rows) > target:
        last = table.rows[-1]._tr
        tbl.remove(last)


def populate_tables(doc, site):
    """Fill the three front-matter tables (now the first three tables)."""
    control, version, refs = doc.tables[0], doc.tables[1], doc.tables[2]

    # Document Control: rows are (label | value). Row 0 is a styled band.
    document_control = site["documentControl"]
    ensure_rows(control, len(document_control), first_data_idx=1)
    for i, (label, value) in enumerate(document_control, start=1):
        row = control.rows[i]
        set_cell_text(row.cells[0], label)
        set_cell_text(row.cells[1], value)

    # Version Control: row 0 is the header.
    version_control = site["versionControl"]
    ensure_rows(version, len(version_control), first_data_idx=1)
    for i, vals in enumerate(version_control, start=1):
        row = version.rows[i]
        for c, val in enumerate(vals):
            set_cell_text(row.cells[c], val)

    # References: row 0 is the header.
    references = site["references"]
    ensure_rows(refs, len(references), first_data_idx=1)
    for i, vals in enumerate(references, start=1):
        row = refs.rows[i]
        for c, val in enumerate(vals):
            set_cell_text(row.cells[c], val)


def patch_chrome_and_props(out_path, chrome_replacements, footer_replacements):
    """Refresh the stale running header + footer text at the zip level (small
    parts; avoids re-serialising document.xml). Header/footer needle->value
    maps come from config; the heading-numbering + TOC-refresh fixes below are
    generic and always applied."""
    with zipfile.ZipFile(out_path) as z:
        parts = {n: z.read(n) for n in z.namelist()}

    if "word/header1.xml" in parts and chrome_replacements:
        h = parts["word/header1.xml"].decode("utf8")
        for needle, value in chrome_replacements.items():
            h = h.replace(needle, value)
        parts["word/header1.xml"] = h.encode("utf8")

    if footer_replacements:
        for fn in ("word/footer1.xml", "word/footer2.xml", "word/footer3.xml",
                   "word/footer4.xml", "word/footer5.xml", "word/footer6.xml"):
            if fn in parts:
                f = parts[fn].decode("utf8")
                for needle, value in footer_replacements.items():
                    f = f.replace(needle, value)
                parts[fn] = f.encode("utf8")

    # Disable the template's heading auto-numbering. Our headings already carry
    # manual numbers ("1.", "2.1", …) that the inline cross-references depend on;
    # the template's Heading styles also auto-number (1. / a. / i.), which would
    # render "1. 1. Executive Summary". Strip <w:numPr> from the Heading styles
    # so only our manual numbers show.
    if "word/styles.xml" in parts:
        styles = parts["word/styles.xml"].decode("utf8")
        def _strip_numpr(m):
            return re.sub(r"<w:numPr>.*?</w:numPr>", "", m.group(0), flags=re.S)
        styles = re.sub(
            r'<w:style\b[^>]*w:styleId="Heading[1-9]"[^>]*>.*?</w:style>',
            _strip_numpr, styles, flags=re.S)
        parts["word/styles.xml"] = styles.encode("utf8")

    # Make Word rebuild the TOC field from our headings when the file opens.
    if "word/settings.xml" in parts:
        st = parts["word/settings.xml"].decode("utf8")
        if "updateFields" not in st:
            st = re.sub(r"(<w:settings\b[^>]*>)",
                        r'\1<w:updateFields w:val="true"/>', st, count=1)
            parts["word/settings.xml"] = st.encode("utf8")

    tmp = out_path + ".tmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in parts.items():
            z.writestr(name, data)
    os.replace(tmp, out_path)


def main():
    here = os.path.dirname(os.path.abspath(__file__))
    repo = os.path.dirname(here)
    default_config = os.environ.get(
        "BRANDED_DOCX_CONFIG", os.path.join(here, "branded_docx.config.json"))

    ap = argparse.ArgumentParser(description="Build a branded DESIGN_DOC.docx.")
    ap.add_argument("--template", default=os.environ.get("BRANDED_TEMPLATE"),
                    help="Path to the branded reference .docx "
                         "(or set BRANDED_TEMPLATE).")
    ap.add_argument("--config", default=default_config,
                    help="Path to the site config JSON "
                         "(or set BRANDED_DOCX_CONFIG).")
    ap.add_argument("--source", default=os.path.join(repo, "docs", "DESIGN_DOC.md"))
    ap.add_argument("--out", default=os.path.join(repo, "docs", "DESIGN_DOC.docx"))
    args = ap.parse_args()

    if not args.template or not os.path.isfile(args.template):
        raise SystemExit(
            "build_branded_docx: branded reference template not found.\n"
            "  Pass --template <path> or set BRANDED_TEMPLATE.\n"
            f"  (got: {args.template!r})")

    site = load_site(args.config)
    meta = site["meta"]

    work = tempfile.mkdtemp(prefix="branded_docx_")
    try:
        # 1. Stripped body markdown.
        with open(args.source, encoding="utf8") as fh:
            body_md = strip_front_matter(fh.read())
        body_md_path = os.path.join(work, "body.md")
        with open(body_md_path, "w", encoding="utf8") as fh:
            fh.write(body_md)

        # 2. Render body in the template's visual system.
        content_path = os.path.join(work, "content.docx")
        run_pandoc(body_md_path, args.template, content_path)

        # 3. Pull cover + control tables + section properties from the template.
        front, body_sect, root_tag = extract_template_parts(args.template)
        front = edit_cover_text(front, site["cover"])
        front_children = parse_fragment(root_tag, front)
        front_children = swap_stale_toc(front_children, root_tag)
        body_sect_el = parse_sectpr(root_tag, body_sect)

        # 4. Splice: prepend front matter, restore template body section props.
        doc = Document(content_path)
        body = doc.element.body
        old_sect = body.find(w("sectPr"))
        if old_sect is not None:
            body.remove(old_sect)
        body.append(body_sect_el)              # body sectPr must stay last
        for i, el in enumerate(front_children):
            body.insert(i, el)                 # cover + tables go first

        # 5. Document text onto the control tables + core properties.
        #    The cover title/subtitle are content controls bound to the core
        #    Title/Subject properties (Word refreshes them from here), so these
        #    drive the cover: Title -> big title, Subject -> subtitle.
        populate_tables(doc, site)
        cp = doc.core_properties
        cp.title = meta["title"]
        cp.subject = meta["subtitle"]
        cp.author = meta["author"]
        cp.keywords = meta["keywords"]

        doc.save(args.out)

        # 6. Refresh stale header/footer text (separate parts).
        patch_chrome_and_props(args.out, site["chrome"], site["footer"])
        print(f"build_branded_docx: wrote {args.out}")
    finally:
        shutil.rmtree(work, ignore_errors=True)


if __name__ == "__main__":
    main()
