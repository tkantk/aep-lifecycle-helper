"""
Convert docs/DESIGN_DOC.md → docs/DESIGN_DOC.docx using python-docx.

The python-docx library doesn't parse Markdown itself; this script walks the
Markdown line-by-line and emits the corresponding Word structures:

  - `# ` / `## ` / `### ` / `#### ` → heading 1–4
  - ``` fenced code blocks → monospace paragraphs with a light-grey shading
  - tables (pipe-delimited) → Word tables with the default grid style
  - `- ` / `* ` lists → bulleted paragraphs
  - `1. ` / `2. ` lists → numbered paragraphs
  - **bold**, *italic*, `inline code` → runs with the matching formatting
  - blank lines → paragraph breaks

Anything more exotic (nested lists, images, blockquotes, HTML in Markdown) is
rendered as plain text — fine for our design doc which avoids those.

Usage:
    python scripts/md-to-docx.py docs/DESIGN_DOC.md docs/DESIGN_DOC.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


CODE_FONT = "Consolas"
CODE_SIZE = 9   # points

BOLD_RE       = re.compile(r"\*\*([^*]+)\*\*")
ITALIC_RE     = re.compile(r"(?<!\*)\*(?!\*)([^*]+)\*(?!\*)")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")
# Combined tokenizer — order matters: bold before italic (both use asterisks)
TOKEN_RE = re.compile(
    r"(\*\*[^*]+\*\*|`[^`]+`|(?<!\*)\*(?!\*)[^*]+\*(?!\*))"
)


def shade_paragraph(paragraph, fill_hex: str) -> None:
    """Add a shading (background color) to a paragraph's XML."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def add_inline_runs(paragraph, text: str) -> None:
    """Tokenize a line for **bold**, *italic*, `code`, emit runs accordingly."""
    pos = 0
    for match in TOKEN_RE.finditer(text):
        start, end = match.span()
        if start > pos:
            paragraph.add_run(text[pos:start])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(CODE_SIZE)
        else:  # italic
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_code_block(doc: Document, code_text: str) -> None:
    """Render a fenced code block as a shaded monospace paragraph."""
    paragraph = doc.add_paragraph()
    shade_paragraph(paragraph, "F3F3F3")
    run = paragraph.add_run(code_text)
    run.font.name = CODE_FONT
    run.font.size = Pt(CODE_SIZE)
    # Keep the lines tight
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.space_before = Pt(6)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    """Render a Markdown-style table (rows already parsed) as a Word table."""
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, cell_text.strip())
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True


def is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    # Default body font a bit tighter for long docs
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block boundary
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Tables — consume consecutive table rows as a single block
        if is_table_row(line):
            table_lines: list[str] = []
            while i < len(lines) and is_table_row(lines[i]):
                table_lines.append(lines[i])
                i += 1
            # Drop the separator row (|---|---|)
            rows = [
                [c.strip() for c in row.strip().strip("|").split("|")]
                for row in table_lines
                if not re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", row)
            ]
            add_table(doc, rows)
            continue

        # Headings
        if line.startswith("#"):
            m = re.match(r"^(#{1,6})\s+(.*)$", line)
            if m:
                level = min(len(m.group(1)), 4)
                heading = doc.add_heading(level=level)
                add_inline_runs(heading, m.group(2))
                i += 1
                continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            # Emit as an empty paragraph with a top border
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), "6")
            top.set(qn("w:space"), "1")
            top.set(qn("w:color"), "CCCCCC")
            pBdr.append(top)
            pPr.append(pBdr)
            i += 1
            continue

        # Bullet / numbered list
        bullet_match = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, bullet_match.group(1))
            i += 1
            continue
        numbered_match = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered_match:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, numbered_match.group(1))
            i += 1
            continue

        # Blank line
        if line.strip() == "":
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, line)
        i += 1

    doc.save(str(docx_path))


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Usage: python scripts/md-to-docx.py <input.md> <output.docx>", file=sys.stderr)
        return 2
    md = Path(argv[1])
    docx = Path(argv[2])
    if not md.exists():
        print(f"Input file not found: {md}", file=sys.stderr)
        return 1
    convert(md, docx)
    print(f"Wrote {docx} ({docx.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
